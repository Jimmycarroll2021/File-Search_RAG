"""
Tests for export service - markdown to PDF/DOCX conversion
"""
import pytest
import io
from datetime import datetime
from docx import Document
from app.services.export_service import ExportService


class TestMarkdownToDOCX:
    """Tests for markdown to DOCX conversion"""

    def test_basic_markdown_to_docx(self):
        """Test basic markdown conversion to DOCX"""
        markdown_text = "# Heading 1\n\nThis is a paragraph with **bold** and *italic* text."
        metadata = {
            "title": "Test Document",
            "question": "What is the answer?",
            "date": "2025-11-09",
            "response_mode": "Tender Response",
            "company": "Test Company"
        }

        service = ExportService()
        docx_buffer = service.markdown_to_docx(markdown_text, metadata)

        # Verify buffer is BytesIO
        assert isinstance(docx_buffer, io.BytesIO)

        # Load document and verify content
        docx_buffer.seek(0)
        doc = Document(docx_buffer)

        # Should have paragraphs (title, question, date, content)
        assert len(doc.paragraphs) > 0

        # Verify title is present
        title_found = any("Test Document" in p.text for p in doc.paragraphs)
        assert title_found, "Title not found in document"

    def test_docx_with_code_blocks(self):
        """Test markdown with code blocks"""
        markdown_text = """
# Code Example

Here is some code:

```python
def hello():
    print("Hello, World!")
```

End of code.
"""
        metadata = {
            "title": "Code Document",
            "question": None,
            "date": "2025-11-09"
        }

        service = ExportService()
        docx_buffer = service.markdown_to_docx(markdown_text, metadata)

        docx_buffer.seek(0)
        doc = Document(docx_buffer)

        # Verify code block is present
        code_found = any("def hello()" in p.text for p in doc.paragraphs)
        assert code_found, "Code block not found in document"

    def test_docx_with_lists(self):
        """Test markdown with lists"""
        markdown_text = """
# List Test

Bulleted list:
- Item 1
- Item 2
- Item 3

Numbered list:
1. First
2. Second
3. Third
"""
        metadata = {"title": "List Document", "date": "2025-11-09"}

        service = ExportService()
        docx_buffer = service.markdown_to_docx(markdown_text, metadata)

        docx_buffer.seek(0)
        doc = Document(docx_buffer)

        # Verify list content exists
        assert len(doc.paragraphs) > 0

    def test_docx_with_table(self):
        """Test markdown with tables"""
        markdown_text = """
# Table Test

| Header 1 | Header 2 | Header 3 |
|----------|----------|----------|
| Cell 1   | Cell 2   | Cell 3   |
| Cell 4   | Cell 5   | Cell 6   |
"""
        metadata = {"title": "Table Document", "date": "2025-11-09"}

        service = ExportService()
        docx_buffer = service.markdown_to_docx(markdown_text, metadata)

        docx_buffer.seek(0)
        doc = Document(docx_buffer)

        # Verify table exists
        assert len(doc.tables) > 0, "No tables found in document"
        table = doc.tables[0]
        assert len(table.rows) == 3, "Expected 3 rows in table"
        assert len(table.columns) == 3, "Expected 3 columns in table"

    def test_docx_metadata_optional_fields(self):
        """Test DOCX generation with optional metadata fields"""
        markdown_text = "# Simple Test"
        metadata = {
            "title": "Minimal Document",
            "date": "2025-11-09"
            # No question, response_mode, or company
        }

        service = ExportService()
        docx_buffer = service.markdown_to_docx(markdown_text, metadata)

        assert isinstance(docx_buffer, io.BytesIO)
        docx_buffer.seek(0)
        doc = Document(docx_buffer)
        assert len(doc.paragraphs) > 0


class TestMarkdownToPDF:
    """Tests for markdown to PDF conversion"""

    def test_basic_markdown_to_pdf(self):
        """Test basic markdown conversion to PDF"""
        markdown_text = "# Heading 1\n\nThis is a paragraph with **bold** and *italic* text."
        metadata = {
            "title": "Test PDF Document",
            "question": "What is the answer?",
            "date": "2025-11-09",
            "response_mode": "Tender Response",
            "company": "Test Company"
        }

        service = ExportService()
        pdf_buffer = service.markdown_to_pdf(markdown_text, metadata)

        # Verify buffer is BytesIO
        assert isinstance(pdf_buffer, io.BytesIO)

        # Verify PDF has content (PDF starts with %PDF)
        pdf_buffer.seek(0)
        pdf_content = pdf_buffer.read(4)
        assert pdf_content == b'%PDF', "Not a valid PDF file"

    def test_pdf_with_code_blocks(self):
        """Test PDF generation with code blocks"""
        markdown_text = """
# Code Example

Here is some code:

```python
def hello():
    print("Hello, World!")
```
"""
        metadata = {"title": "Code PDF", "date": "2025-11-09"}

        service = ExportService()
        pdf_buffer = service.markdown_to_pdf(markdown_text, metadata)

        pdf_buffer.seek(0)
        pdf_content = pdf_buffer.read(4)
        assert pdf_content == b'%PDF', "Not a valid PDF file"

    def test_pdf_with_table(self):
        """Test PDF generation with tables"""
        markdown_text = """
# Table Test

| Header 1 | Header 2 |
|----------|----------|
| Cell 1   | Cell 2   |
"""
        metadata = {"title": "Table PDF", "date": "2025-11-09"}

        service = ExportService()
        pdf_buffer = service.markdown_to_pdf(markdown_text, metadata)

        pdf_buffer.seek(0)
        pdf_content = pdf_buffer.read(4)
        assert pdf_content == b'%PDF', "Not a valid PDF file"

    def test_pdf_metadata_optional_fields(self):
        """Test PDF generation with minimal metadata"""
        markdown_text = "# Simple Test"
        metadata = {
            "title": "Minimal PDF",
            "date": "2025-11-09"
        }

        service = ExportService()
        pdf_buffer = service.markdown_to_pdf(markdown_text, metadata)

        assert isinstance(pdf_buffer, io.BytesIO)
        pdf_buffer.seek(0)
        pdf_content = pdf_buffer.read(4)
        assert pdf_content == b'%PDF', "Not a valid PDF file"


class TestExportServiceHelpers:
    """Tests for helper functions in export service"""

    def test_markdown_to_html_conversion(self):
        """Test markdown to HTML conversion helper"""
        service = ExportService()
        markdown_text = "# Header\n\n**Bold** text"

        html = service._markdown_to_html(markdown_text)

        assert "<h1>Header</h1>" in html
        assert "<strong>Bold</strong>" in html

    def test_sanitize_filename(self):
        """Test filename sanitization"""
        service = ExportService()

        # Test with invalid characters
        result = service._sanitize_filename("My/File\\Name:Test*.pdf")
        assert "/" not in result
        assert "\\" not in result
        assert ":" not in result
        assert "*" not in result

        # Test with valid filename
        result = service._sanitize_filename("MyFile123.pdf")
        assert result == "MyFile123.pdf"
