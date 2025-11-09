"""
Export service - Convert markdown to PDF and DOCX formats
"""
import io
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import markdown
from xhtml2pdf import pisa


class ExportService:
    """Service for exporting markdown content to PDF and DOCX formats"""

    def __init__(self):
        """Initialize the export service"""
        pass

    def markdown_to_docx(self, markdown_text, metadata):
        """
        Convert markdown text to a styled DOCX document.

        Args:
            markdown_text: Markdown formatted text
            metadata: Dict containing title, question, date, response_mode, company

        Returns:
            BytesIO buffer containing the DOCX document
        """
        doc = Document()

        # Add custom styles
        self._setup_docx_styles(doc)

        # Add metadata header
        self._add_docx_header(doc, metadata)

        # Parse and add markdown content
        self._parse_markdown_to_docx(doc, markdown_text)

        # Add footer with page numbers
        self._add_docx_footer(doc)

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    def markdown_to_pdf(self, markdown_text, metadata):
        """
        Convert markdown text to PDF via HTML template.

        Args:
            markdown_text: Markdown formatted text
            metadata: Dict containing title, question, date, response_mode, company

        Returns:
            BytesIO buffer containing the PDF document
        """
        # Convert markdown to HTML
        html_content = self._markdown_to_html(markdown_text)

        # Build HTML with metadata
        rendered_html = self._build_pdf_html(
            title=metadata.get('title', 'Document'),
            question=metadata.get('question'),
            date=metadata.get('date', datetime.now().strftime('%Y-%m-%d')),
            response_mode=metadata.get('response_mode'),
            company=metadata.get('company'),
            content=html_content
        )

        # Generate PDF using xhtml2pdf
        pdf_buffer = io.BytesIO()
        pisa_status = pisa.CreatePDF(
            rendered_html,
            dest=pdf_buffer
        )

        if pisa_status.err:
            raise Exception("Error generating PDF")

        pdf_buffer.seek(0)
        return pdf_buffer

    def _markdown_to_html(self, markdown_text):
        """Convert markdown to HTML"""
        return markdown.markdown(
            markdown_text,
            extensions=['tables', 'fenced_code', 'nl2br']
        )

    def _sanitize_filename(self, filename):
        """Sanitize filename by removing invalid characters"""
        # Remove invalid characters
        sanitized = re.sub(r'[<>:"/\\|?*]', '', filename)
        # Remove leading/trailing spaces and dots
        sanitized = sanitized.strip('. ')
        return sanitized

    def _setup_docx_styles(self, doc):
        """Setup custom styles for DOCX document"""
        styles = doc.styles

        # Title style
        if 'DocTitle' not in styles:
            title_style = styles.add_style('DocTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_format = title_style.font
            title_format.name = 'Calibri'
            title_format.size = Pt(18)
            title_format.bold = True
            title_format.color.rgb = RGBColor(0, 0, 0)

        # Question style
        if 'Question' not in styles:
            question_style = styles.add_style('Question', WD_STYLE_TYPE.PARAGRAPH)
            question_format = question_style.font
            question_format.name = 'Calibri'
            question_format.size = Pt(12)
            question_format.italic = True
            question_format.color.rgb = RGBColor(128, 128, 128)

        # Code style
        if 'CodeBlock' not in styles:
            code_style = styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
            code_format = code_style.font
            code_format.name = 'Consolas'
            code_format.size = Pt(10)
            code_style.paragraph_format.left_indent = Inches(0.5)
            code_style.paragraph_format.space_before = Pt(6)
            code_style.paragraph_format.space_after = Pt(6)

    def _add_docx_header(self, doc, metadata):
        """Add header with metadata to DOCX"""
        # Title
        title_para = doc.add_paragraph(metadata.get('title', 'Document'))
        title_para.style = 'DocTitle'

        # Question (if provided)
        if metadata.get('question'):
            question_para = doc.add_paragraph(f"Question: {metadata['question']}")
            question_para.style = 'Question'

        # Metadata line
        meta_parts = []
        if metadata.get('date'):
            meta_parts.append(f"Date: {metadata['date']}")
        if metadata.get('response_mode'):
            meta_parts.append(f"Mode: {metadata['response_mode']}")
        if metadata.get('company'):
            meta_parts.append(f"Company: {metadata['company']}")

        if meta_parts:
            meta_para = doc.add_paragraph(' | '.join(meta_parts))
            meta_para.style = 'Question'

        # Separator
        doc.add_paragraph()

    def _add_docx_footer(self, doc):
        """Add footer with page numbers to DOCX"""
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "Page "
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _parse_markdown_to_docx(self, doc, markdown_text):
        """Parse markdown and add to DOCX with proper formatting"""
        lines = markdown_text.split('\n')
        i = 0
        in_code_block = False
        in_table = False
        table_rows = []
        code_block_lines = []

        while i < len(lines):
            line = lines[i]

            # Handle code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    # End of code block
                    code_text = '\n'.join(code_block_lines)
                    code_para = doc.add_paragraph(code_text)
                    code_para.style = 'CodeBlock'
                    # Add gray background (approximation)
                    shading = code_para._element.get_or_add_pPr()
                    code_block_lines = []
                    in_code_block = False
                else:
                    # Start of code block
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_block_lines.append(line)
                i += 1
                continue

            # Handle tables
            if '|' in line and line.strip().startswith('|'):
                if not in_table:
                    in_table = True
                    table_rows = []
                table_rows.append(line)
                i += 1
                continue
            elif in_table:
                # End of table
                self._add_table_to_docx(doc, table_rows)
                table_rows = []
                in_table = False

            # Handle headings
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                if level == 1:
                    para = doc.add_heading(text, level=1)
                elif level == 2:
                    para = doc.add_heading(text, level=2)
                elif level == 3:
                    para = doc.add_heading(text, level=3)
                else:
                    para = doc.add_heading(text, level=4)
                i += 1
                continue

            # Handle lists
            if line.strip().startswith(('- ', '* ', '+ ')):
                text = line.strip()[2:]
                text = self._apply_inline_formatting(text)
                doc.add_paragraph(text, style='List Bullet')
                i += 1
                continue

            if re.match(r'^\d+\.\s', line.strip()):
                text = re.sub(r'^\d+\.\s', '', line.strip())
                text = self._apply_inline_formatting(text)
                doc.add_paragraph(text, style='List Number')
                i += 1
                continue

            # Regular paragraph
            if line.strip():
                text = self._apply_inline_formatting(line)
                doc.add_paragraph(text)

            i += 1

        # Handle table at end of document
        if in_table and table_rows:
            self._add_table_to_docx(doc, table_rows)

    def _add_table_to_docx(self, doc, table_rows):
        """Add a markdown table to DOCX"""
        # Parse table rows
        rows = []
        for row in table_rows:
            cells = [cell.strip() for cell in row.split('|')[1:-1]]
            if cells and not all(cell.startswith('-') for cell in cells):
                rows.append(cells)

        if not rows:
            return

        # Create table
        num_cols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'

        # Populate table
        for i, row_data in enumerate(rows):
            for j, cell_data in enumerate(row_data):
                table.rows[i].cells[j].text = cell_data
                # Bold header row
                if i == 0:
                    for paragraph in table.rows[i].cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

    def _apply_inline_formatting(self, text):
        """Apply bold and italic formatting to text (simplified)"""
        # Note: This returns plain text with markdown removed
        # For actual formatting, we'd need to work with runs
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Bold
        text = re.sub(r'\*(.+?)\*', r'\1', text)  # Italic
        text = re.sub(r'`(.+?)`', r'\1', text)  # Inline code
        return text

    def _build_pdf_html(self, title, question, date, response_mode, company, content):
        """Build HTML for PDF generation with metadata"""
        question_html = f'<p class="question">Question: {question}</p>' if question else ''

        metadata_parts = []
        if date:
            metadata_parts.append(f'<span>Date: {date}</span>')
        if response_mode:
            metadata_parts.append(f'<span> | Mode: {response_mode}</span>')
        if company:
            metadata_parts.append(f'<span> | Company: {company}</span>')
        metadata_html = ''.join(metadata_parts)

        # Use regular string concatenation to avoid f-string curly brace issues
        html = '''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>''' + title + '''</title>
    <style>
        @page {
            size: A4;
            margin: 1in;
        }
        body {
            font-family: Arial, Helvetica, sans-serif;
            font-size: 11pt;
            line-height: 1.6;
            color: #333;
        }
        .header {
            margin-bottom: 30px;
            border-bottom: 2px solid #333;
            padding-bottom: 15px;
        }
        h1 {
            font-size: 18pt;
            font-weight: bold;
            margin: 0 0 10px 0;
            color: #000;
        }
        .question {
            font-style: italic;
            color: #666;
            font-size: 12pt;
            margin: 5px 0;
        }
        .metadata {
            font-size: 10pt;
            color: #666;
            margin-top: 5px;
        }
        .content h1 {
            font-size: 14pt;
            font-weight: bold;
            margin-top: 20px;
            margin-bottom: 10px;
        }
        .content h2 {
            font-size: 12pt;
            font-weight: bold;
            margin-top: 15px;
            margin-bottom: 8px;
        }
        .content h3 {
            font-size: 11pt;
            font-weight: bold;
            margin-top: 12px;
            margin-bottom: 6px;
        }
        p {
            margin: 10px 0;
        }
        code {
            font-family: "Courier New", Consolas, monospace;
            font-size: 10pt;
            background-color: #f4f4f4;
            padding: 2px 5px;
        }
        pre {
            font-family: "Courier New", Consolas, monospace;
            font-size: 10pt;
            background-color: #f4f4f4;
            padding: 15px;
            border-left: 4px solid #ccc;
            overflow-x: auto;
            margin: 15px 0;
        }
        pre code {
            background-color: transparent;
            padding: 0;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }
        th, td {
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }
        th {
            background-color: #f4f4f4;
            font-weight: bold;
        }
        ul, ol {
            margin: 10px 0;
            padding-left: 30px;
        }
        li {
            margin: 5px 0;
        }
        strong {
            font-weight: bold;
        }
        em {
            font-style: italic;
        }
        a {
            color: #0066cc;
            text-decoration: underline;
        }
    </style>
</head>
<body>
    <div class="header">
        <h1>''' + title + '''</h1>
        ''' + question_html + '''
        <div class="metadata">
            ''' + metadata_html + '''
        </div>
    </div>
    <div class="content">
        ''' + content + '''
    </div>
</body>
</html>'''
        return html
